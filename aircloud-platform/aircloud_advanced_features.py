#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
🔧 AIRCLOUD LEGAL PLATFORM - ADVANCED FEATURES EXTENSION
========================================================
Rozszerzenie z zaawansowanymi funkcjami dla platformy prawnej

🎯 Nowe funkcjonalności:
• Real-time collaboration (WebSocket) - [OPTIONAL]
• Advanced document versioning
• Legal template system
• Document export (PDF, DOCX)
• Email notifications
• Advanced analytics
• Workflow automation
• Integration APIs

Author: Łukasz Kołodziej | Aircloud
Version: 2.1.0 (Extended)
"""

from flask import Flask, request, jsonify, session, redirect, url_for, flash, send_file, abort
# Real-time features temporarily disabled
# from flask_socketio import SocketIO, emit, join_room, leave_room
import threading
import time
from datetime import datetime, timedelta
import json
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib import colors
import markdown
from docx import Document as DocxDocument
from docx.shared import Inches
import zipfile
import csv

class AircloudAdvancedFeatures:
    """Zaawansowane funkcje Aircloud Legal Platform"""
    
    @staticmethod
    def export_document_to_pdf(document):
        """Eksport dokumentu do PDF"""
        buffer = io.BytesIO()
        doc = SimpleDocDocument(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Custom style for legal documents
        legal_style = ParagraphStyle(
            'LegalDocument',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leftIndent=20,
            fontName='Times-Roman'
        )
        
        title_style = ParagraphStyle(
            'LegalTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            alignment=1,  # Center
            fontName='Times-Bold'
        )
        
        story = []
        
        # Header with Aircloud branding
        header_text = f"""
        <para align="center">
        <b>AIRCLOUD LEGAL PLATFORM</b><br/>
        Professional Document System<br/>
        © 2025 Łukasz Kołodziej | Aircloud
        </para>
        """
        story.append(Paragraph(header_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Document title
        story.append(Paragraph(document.title, title_style))
        story.append(Spacer(1, 12))
        
        # Document metadata
        metadata = f"""
        <b>System:</b> {document.legal_system.name}<br/>
        <b>Typ:</b> {document.document_type}<br/>
        <b>Autor:</b> {document.author.full_name}<br/>
        <b>Data:</b> {document.created_at.strftime('%d.%m.%Y')}<br/>
        <b>Status:</b> {document.status}<br/>
        """
        story.append(Paragraph(metadata, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Document content (convert markdown to HTML first)
        if document.content_markdown:
            # Convert markdown to HTML then to reportlab
            html_content = markdown.markdown(document.content_markdown)
            # Simple markdown conversion for PDF
            paragraphs = document.content_markdown.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.startswith('#'):
                        # Handle headers
                        level = len(para) - len(para.lstrip('#'))
                        text = para.lstrip('#').strip()
                        if level <= 2:
                            story.append(Paragraph(f"<b>{text}</b>", styles['Heading2']))
                    else:
                        story.append(Paragraph(para, legal_style))
                    story.append(Spacer(1, 6))
        
        # Footer
        footer_text = f"""
        <para align="center">
        Dokument wygenerowany przez Aircloud Legal Platform<br/>
        {datetime.now().strftime('%d.%m.%Y %H:%M')}<br/>
        www.aircloud.pl | legal@aircloud.pl
        </para>
        """
        story.append(Spacer(1, 30))
        story.append(Paragraph(footer_text, styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def export_document_to_docx(document):
        """Eksport dokumentu do DOCX"""
        doc = DocxDocument()
        
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "AIRCLOUD LEGAL PLATFORM - Professional Document System"
        
        # Title
        title = doc.add_heading(document.title, 0)
        title.alignment = 1  # Center
        
        # Metadata table
        doc.add_heading('Informacje o dokumencie', level=2)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Shading Accent 1'
        
        metadata_items = [
            ('System prawny:', document.legal_system.name),
            ('Typ dokumentu:', document.document_type),
            ('Autor:', document.author.full_name),
            ('Data utworzenia:', document.created_at.strftime('%d.%m.%Y')),
            ('Status:', document.status)
        ]
        
        for i, (key, value) in enumerate(metadata_items):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
        
        doc.add_page_break()
        
        # Content
        doc.add_heading('Treść dokumentu', level=2)
        if document.content_markdown:
            # Simple markdown to docx conversion
            paragraphs = document.content_markdown.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.startswith('#'):
                        level = len(para) - len(para.lstrip('#'))
                        text = para.lstrip('#').strip()
                        doc.add_heading(text, level=min(level, 3))
                    else:
                        doc.add_paragraph(para)
        
        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"© 2025 Aircloud Legal Platform | {datetime.now().strftime('%d.%m.%Y')}"
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def generate_legal_template(template_type, system_name="", organization=""):
        """Generuje szablon dokumentu prawnego"""
        
        templates = {
            'regulamin': f"""# REGULAMIN {organization.upper()}

*Wygenerowany przez Aircloud Legal Platform*
*© 2025 Łukasz Kołodziej | Aircloud*

## Rozdział I - Postanowienia ogólne

### § 1. Podstawowe definicje
1. **{organization}** - organizacja działająca zgodnie z niniejszym regulaminem
2. **Członek** - osoba należąca do {organization}
3. **Zarząd** - organ wykonawczy {organization}

### § 2. Zakres stosowania
Niniejszy regulamin określa:
1) Cele i zadania {organization}
2) Strukturę organizacyjną
3) Zasady członkostwa
4) Procedury podejmowania decyzji

## Rozdział II - Cele i zadania

### § 3. Cele statutowe
{organization} działa w celu:
1) [Cel podstawowy - do określenia]
2) [Cel dodatkowy - do określenia]
3) [Inne cele - do określenia]

## Rozdział III - Członkostwo

### § 4. Nabywanie członkostwa
1. Członkiem {organization} może zostać osoba, która:
   - spełnia warunki określone w statucie
   - złoży pisemną deklarację członkowską
   - zostanie przyjęta przez Zarząd

### § 5. Prawa członków
Członek {organization} ma prawo do:
1) Uczestniczenia w zebraniach
2) Wybierania i bycia wybieranym do organów
3) Składania wniosków i postulatów

### § 6. Obowiązki członków
Członek {organization} jest obowiązany:
1) Przestrzegać postanowień regulaminu
2) Realizować podjęte zobowiązania
3) Dbać o dobre imię organizacji

## Rozdział IV - Organy

### § 7. Organy {organization}
Organami {organization} są:
1) **Zebranie Ogólne** - najwyższy organ
2) **Zarząd** - organ wykonawczy
3) **Komisja Rewizyjna** - organ kontrolny

## Rozdział V - Postanowienia końcowe

### § 8. Zmiany regulaminu
1. Zmiany regulaminu mogą być dokonane przez Zebranie Ogólne
2. Zmiana wymaga większości 2/3 głosów obecnych

### § 9. Wejście w życie
Regulamin wchodzi w życie z dniem uchwalenia.

---
**Regulamin uchwalony:** [data]
**Przez:** [organ uchwalający]

*Dokument wygenerowany w systemie Aircloud Legal Platform*
*Professional Legal Collaboration System*
*Autor: Łukasz Kołodziej | © 2025 Aircloud*
""",

            'umowa': f"""# UMOWA O WSPÓŁPRACY

*Szablon wygenerowany przez Aircloud Legal Platform*
*© 2025 Łukasz Kołodziej | Aircloud*

**Zawarta w dniu:** _________________ 

**Pomiędzy:**

**STRONA 1:**
- Nazwa: {organization}
- Adres: _______________________
- Reprezentowana przez: _________

**STRONA 2:**  
- Nazwa: _______________________
- Adres: _______________________
- Reprezentowana przez: _________

## § 1. PRZEDMIOT UMOWY

1. Przedmiotem umowy jest współpraca w zakresie:
   - [określić zakres współpracy]
   - [dodatkowe ustalenia]

## § 2. OBOWIĄZKI STRON

**Strona 1 zobowiązuje się do:**
1) [obowiązek 1]
2) [obowiązek 2]

**Strona 2 zobowiązuje się do:**
1) [obowiązek 1]
2) [obowiązek 2]

## § 3. WYNAGRODZENIE

1. Za wykonanie przedmiotu umowy przysługuje wynagrodzenie w wysokości: _______ PLN
2. Płatność następuje w terminie: _______

## § 4. OKRES OBOWIĄZYWANIA

Umowa obowiązuje od dnia _______ do dnia _______

## § 5. POSTANOWIENIA KOŃCOWE

1. Wszelkie zmiany umowy wymagają formy pisemnej
2. W sprawach nieuregulowanych stosuje się przepisy Kodeksu Cywilnego
3. Umowa została sporządzona w dwóch jednobrzmiących egzemplarzach

**PODPISY:**

Strona 1: _________________    Strona 2: _________________

*Umowa wygenerowana w Aircloud Legal Platform*
*Professional Legal Document System | www.aircloud.pl*
""",

            'procedura': f"""# PROCEDURA: [NAZWA PROCEDURY]

*Utworzona w Aircloud Legal Platform*
*© 2025 Łukasz Kołodziej | Aircloud*

## 1. CEL PROCEDURY

Niniejsza procedura określa sposób [opisać cel procedury] w {organization}.

## 2. ZAKRES STOSOWANIA

Procedura dotyczy:
- [zakres 1]
- [zakres 2]
- [zakres 3]

## 3. ODPOWIEDZIALNOŚCI

| Rola | Odpowiedzialność |
|------|------------------|
| [Rola 1] | [Opis odpowiedzialności] |
| [Rola 2] | [Opis odpowiedzialności] |

## 4. PRZEBIEG PROCEDURY

### Krok 1: [Nazwa kroku]
**Odpowiedzialny:** [Kto]
**Czas:** [Ile czasu]
**Opis:** [Co należy zrobić]

### Krok 2: [Nazwa kroku]
**Odpowiedzialny:** [Kto]
**Czas:** [Ile czasu]  
**Opis:** [Co należy zrobić]

### Krok 3: [Nazwa kroku]
**Odpowiedzialny:** [Kto]
**Czas:** [Ile czasu]
**Opis:** [Co należy zrobić]

## 5. DOKUMENTY POWIĄZANE

- [Dokument 1]
- [Dokument 2]

## 6. ZAŁĄCZNIKI

- Załącznik 1: [Nazwa]
- Załącznik 2: [Nazwa]

## 7. HISTORIA ZMIAN

| Data | Wersja | Opis zmiany | Autor |
|------|--------|-------------|-------|
| {datetime.now().strftime('%d.%m.%Y')} | 1.0 | Pierwsza wersja | [Autor] |

---
*Procedura utworzona w systemie Aircloud Legal Platform*
*Professional Legal Workflow Management*
*www.aircloud.pl | legal@aircloud.pl*
"""
        }
        
        return templates.get(template_type, "Nieznany typ szablonu")
    
    @staticmethod
    def analyze_document_compliance(document):
        """Analiza zgodności dokumentu z przepisami"""
        compliance_score = 85  # Przykładowa ocena
        
        issues = []
        suggestions = []
        
        # Przykładowa analiza
        content = document.content_markdown.lower() if document.content_markdown else ""
        
        if 'rodo' not in content and 'gdpr' not in content:
            issues.append("Brak odniesienia do RODO/GDPR")
            suggestions.append("Dodaj klauzulę informacyjną RODO")
        
        if 'konstytucja' not in content and document.document_type == 'statute':
            suggestions.append("Rozważ odniesienie do Konstytucji RP")
        
        if len(content) < 500:
            issues.append("Dokument może być zbyt krótki")
            suggestions.append("Rozważ rozszerzenie treści")
        
        return {
            'compliance_score': compliance_score,
            'issues': issues,
            'suggestions': suggestions,
            'analysis_date': datetime.now().isoformat()
        }
    
    @staticmethod
    def generate_document_analytics(document):
        """Generuje analitykę dokumentu"""
        content = document.content_markdown or ""
        
        analytics = {
            'word_count': len(content.split()),
            'character_count': len(content),
            'paragraph_count': len([p for p in content.split('\n\n') if p.strip()]),
            'readability_score': AircloudAdvancedFeatures._calculate_readability(content),
            'complexity_level': AircloudAdvancedFeatures._assess_complexity(content),
            'legal_terms_count': AircloudAdvancedFeatures._count_legal_terms(content),
            'structure_score': AircloudAdvancedFeatures._analyze_structure(content)
        }
        
        return analytics
    
    @staticmethod
    def _calculate_readability(text):
        """Oblicza wskaźnik czytelności"""
        words = text.split()
        sentences = text.count('.') + text.count('!') + text.count('?')
        if sentences == 0:
            return 50
        avg_words_per_sentence = len(words) / sentences
        
        # Uproszczony wskaźnik czytelności
        if avg_words_per_sentence < 15:
            return 90
        elif avg_words_per_sentence < 20:
            return 75
        elif avg_words_per_sentence < 25:
            return 60
        else:
            return 45
    
    @staticmethod
    def _assess_complexity(text):
        """Ocenia złożoność tekstu"""
        legal_complexity_words = [
            'jednakże', 'ponadto', 'w szczególności', 'z zastrzeżeniem',
            'w rozumieniu', 'stosownie', 'zgodnie z', 'z uwagi na'
        ]
        
        complexity_count = sum(text.lower().count(word) for word in legal_complexity_words)
        text_length = len(text.split())
        
        if text_length == 0:
            return 'low'
        
        complexity_ratio = complexity_count / text_length * 100
        
        if complexity_ratio > 5:
            return 'high'
        elif complexity_ratio > 2:
            return 'medium'
        else:
            return 'low'
    
    @staticmethod
    def _count_legal_terms(text):
        """Liczy terminy prawne"""
        legal_terms = [
            'ustawa', 'rozporządzenie', 'konstytucja', 'kodeks',
            'paragraf', 'artykuł', 'przepis', 'norma', 'prawo',
            'sąd', 'wyrok', 'orzeczenie', 'postępowanie'
        ]
        
        return sum(text.lower().count(term) for term in legal_terms)
    
    @staticmethod
    def _analyze_structure(text):
        """Analizuje strukturę dokumentu"""
        has_headers = '#' in text
        has_paragraphs = '§' in text
        has_numbering = any(f'{i}.' in text for i in range(1, 10))
        
        score = 0
        if has_headers:
            score += 30
        if has_paragraphs:
            score += 40
        if has_numbering:
            score += 30
        
        return min(score, 100)

class AircloudNotificationSystem:
    """System powiadomień dla platformy"""
    
    @staticmethod
    def send_email_notification(user_email, subject, content):
        """Wysyła powiadomienie email (placeholder)"""
        # W produkcji tutaj byłaby integracja z SMTP
        print(f"EMAIL NOTIFICATION:")
        print(f"To: {user_email}")
        print(f"Subject: {subject}")
        print(f"Content: {content[:100]}...")
        return True
    
    @staticmethod
    def create_notification(user_id, message, notification_type="info"):
        """Tworzy powiadomienie w systemie"""
        # Tutaj byłaby integracja z bazą powiadomień
        notification = {
            'user_id': user_id,
            'message': message,
            'type': notification_type,
            'created_at': datetime.now().isoformat(),
            'read': False
        }
        return notification

# WebSocket support for real-time collaboration
class AircloudRealTimeCollaboration:
    """Współpraca w czasie rzeczywistym"""
    
    def __init__(self, socketio):
        self.socketio = socketio
        self.active_sessions = {}
    
    def handle_join_document(self, data):
        """Użytkownik dołącza do dokumentu"""
        document_id = data['document_id']
        user_name = data['user_name']
        
        join_room(f"document_{document_id}")
        
        if document_id not in self.active_sessions:
            self.active_sessions[document_id] = []
        
        self.active_sessions[document_id].append({
            'user_name': user_name,
            'joined_at': datetime.now().isoformat()
        })
        
        emit('user_joined', {
            'user_name': user_name,
            'active_users': self.active_sessions[document_id]
        }, room=f"document_{document_id}")
    
    def handle_document_change(self, data):
        """Obsługuje zmiany w dokumencie"""
        document_id = data['document_id']
        change = data['change']
        user_name = data['user_name']
        
        emit('document_updated', {
            'change': change,
            'user_name': user_name,
            'timestamp': datetime.now().isoformat()
        }, room=f"document_{document_id}", include_self=False)
    
    def handle_leave_document(self, data):
        """Użytkownik opuszcza dokument"""
        document_id = data['document_id']
        user_name = data['user_name']
        
        leave_room(f"document_{document_id}")
        
        if document_id in self.active_sessions:
            self.active_sessions[document_id] = [
                user for user in self.active_sessions[document_id] 
                if user['user_name'] != user_name
            ]
        
        emit('user_left', {
            'user_name': user_name,
            'active_users': self.active_sessions.get(document_id, [])
        }, room=f"document_{document_id}")

print("✅ Aircloud Advanced Features Module Loaded Successfully!")
print("🚀 Extended functionality ready for Aircloud Legal Platform")
print("© 2025 Łukasz Kołodziej | Aircloud Professional")