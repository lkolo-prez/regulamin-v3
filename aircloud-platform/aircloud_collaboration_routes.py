#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
🌐 AIRCLOUD COLLABORATION ROUTES
==================================
Flask routes for collaboration features

🔧 Author: Łukasz Kołodziej
🏢 Company: Aircloud
📅 Version: 3.0.0
"""

from flask import Blueprint, render_template, request, jsonify, redirect, url_for, flash
from flask_login import login_required, current_user
from datetime import datetime
import json

from aircloud_collaboration_engine import (
    collab_sessions, 
    DocumentTemplate, 
    WorkflowEngine,
    CommentingSystem,
    DocumentRelationships,
    NotificationSystem
)

# =====================================
# BLUEPRINT INITIALIZATION
# =====================================

collab_routes = Blueprint('collab_routes', __name__, url_prefix='/collaboration')

# =====================================
# COLLABORATIVE EDITING ROUTES
# =====================================

@collab_routes.route('/editor/<int:document_id>')
@login_required
def collaborative_editor(document_id):
    """Collaborative document editor view"""
    from aircloud_legal_platform import Document
    
    document = Document.query.get_or_404(document_id)
    
    # Join collaboration session
    session_data = collab_sessions.join_session(
        document_id, 
        current_user.id, 
        current_user.username
    )
    
    return render_template(
        'collaboration/editor.html',
        document=document,
        active_users=session_data['users'],
        user_color=session_data['users'][current_user.id]['color']
    )

@collab_routes.route('/session/<int:document_id>/users')
@login_required
def get_active_users(document_id):
    """Get list of active users in document session"""
    if document_id in collab_sessions.active_sessions:
        users = collab_sessions.active_sessions[document_id]['users']
        return jsonify({
            'success': True,
            'users': [
                {
                    'id': user_id,
                    'username': data['username'],
                    'color': data['color'],
                    'joined_at': data['joined_at']
                }
                for user_id, data in users.items()
            ]
        })
    return jsonify({'success': True, 'users': []})

@collab_routes.route('/session/<int:document_id>/leave', methods=['POST'])
@login_required
def leave_session(document_id):
    """Leave collaborative session"""
    collab_sessions.leave_session(document_id, current_user.id)
    return jsonify({'success': True})

# =====================================
# DOCUMENT TEMPLATE ROUTES
# =====================================

@collab_routes.route('/templates')
@login_required
def template_gallery():
    """Display template gallery"""
    templates = DocumentTemplate.get_templates()
    return render_template(
        'collaboration/templates.html',
        templates=templates,
        categories=DocumentTemplate.TEMPLATE_CATEGORIES
    )

@collab_routes.route('/templates/<template_id>')
@login_required
def template_details(template_id):
    """Show template details and preview"""
    templates = DocumentTemplate.get_templates()
    
    # Find template
    template = None
    for category_templates in templates.values():
        for t in category_templates:
            if t['id'] == template_id:
                template = t
                break
    
    if not template:
        flash('Szablon nie został znaleziony', 'error')
        return redirect(url_for('collab_routes.template_gallery'))
    
    return render_template(
        'collaboration/template_details.html',
        template=template
    )

@collab_routes.route('/templates/<template_id>/create', methods=['GET', 'POST'])
@login_required
def create_from_template(template_id):
    """Create new document from template"""
    from aircloud_legal_platform import Document, db
    
    if request.method == 'POST':
        data = request.form.to_dict()
        
        # Generate content from template
        content = DocumentTemplate.create_from_template(template_id, data)
        
        # Create document
        document = Document(
            title=data.get('document_title', 'Nowy dokument'),
            content=content,
            author_id=current_user.id,
            template_id=template_id,
            status='draft',
            created_at=datetime.utcnow()
        )
        
        db.session.add(document)
        db.session.commit()
        
        flash(f'Dokument "{document.title}" został utworzony z szablonu!', 'success')
        return redirect(url_for('collab_routes.collaborative_editor', document_id=document.id))
    
    # GET request - show form
    templates = DocumentTemplate.get_templates()
    template = None
    for category_templates in templates.values():
        for t in category_templates:
            if t['id'] == template_id:
                template = t
                break
    
    return render_template(
        'collaboration/create_from_template.html',
        template=template
    )

# =====================================
# WORKFLOW MANAGEMENT ROUTES
# =====================================

@collab_routes.route('/workflow/<int:document_id>')
@login_required
def workflow_view(document_id):
    """View document workflow status"""
    from aircloud_legal_platform import Document
    
    document = Document.query.get_or_404(document_id)
    
    # Get available actions for user
    user_role = current_user.role if hasattr(current_user, 'role') else 'member'
    available_actions = WorkflowEngine.get_available_actions(
        document.status, 
        user_role
    )
    
    return render_template(
        'collaboration/workflow.html',
        document=document,
        workflow_states=WorkflowEngine.WORKFLOW_STATES,
        available_actions=available_actions,
        workflow_actions=WorkflowEngine.WORKFLOW_ACTIONS
    )

@collab_routes.route('/workflow/<int:document_id>/action', methods=['POST'])
@login_required
def execute_workflow_action(document_id):
    """Execute workflow action"""
    from aircloud_legal_platform import Document, db
    
    action = request.form.get('action')
    comment = request.form.get('comment', '')
    
    success = WorkflowEngine.execute_action(
        document_id,
        action,
        current_user.id,
        comment
    )
    
    if success:
        # Update document status
        document = Document.query.get(document_id)
        # Map action to new status (simplified)
        status_map = {
            'submit_for_review': 'review',
            'approve': 'approved',
            'reject': 'rejected',
            'request_changes': 'draft'
        }
        if action in status_map:
            document.status = status_map[action]
            document.updated_at = datetime.utcnow()
            db.session.commit()
        
        flash('Akcja wykonana pomyślnie!', 'success')
    else:
        flash('Nie udało się wykonać akcji', 'error')
    
    return redirect(url_for('collab_routes.workflow_view', document_id=document_id))

# =====================================
# ADVANCED COMMENTING ROUTES
# =====================================

@collab_routes.route('/comments/<int:document_id>/add', methods=['POST'])
@login_required
def add_contextual_comment(document_id):
    """Add contextual comment to document"""
    data = request.get_json()
    
    comment = CommentingSystem.add_contextual_comment(
        document_id=document_id,
        user_id=current_user.id,
        content=data.get('content'),
        context=data.get('context', {})
    )
    
    # In real implementation, save to database
    
    return jsonify({
        'success': True,
        'comment': comment
    })

@collab_routes.route('/comments/<int:document_id>/paragraph/<paragraph_id>')
@login_required
def get_paragraph_comments(document_id, paragraph_id):
    """Get comments for specific paragraph"""
    comments = CommentingSystem.get_comments_for_paragraph(document_id, paragraph_id)
    
    return jsonify({
        'success': True,
        'comments': comments
    })

@collab_routes.route('/comments/<int:comment_id>/resolve', methods=['POST'])
@login_required
def resolve_comment(comment_id):
    """Mark comment as resolved"""
    success = CommentingSystem.resolve_comment(comment_id, current_user.id)
    
    return jsonify({
        'success': success,
        'message': 'Komentarz oznaczony jako rozwiązany' if success else 'Błąd'
    })

# =====================================
# DOCUMENT RELATIONSHIPS ROUTES
# =====================================

@collab_routes.route('/relationships/<int:document_id>')
@login_required
def view_relationships(document_id):
    """View document relationships"""
    from aircloud_legal_platform import Document
    
    document = Document.query.get_or_404(document_id)
    relationships = DocumentRelationships.get_related_documents(document_id)
    
    return render_template(
        'collaboration/relationships.html',
        document=document,
        relationships=relationships,
        relationship_types=DocumentRelationships.RELATIONSHIP_TYPES
    )

@collab_routes.route('/relationships/<int:source_id>/add', methods=['POST'])
@login_required
def add_relationship(source_id):
    """Add relationship between documents"""
    target_id = request.form.get('target_document_id', type=int)
    relationship_type = request.form.get('relationship_type')
    description = request.form.get('description', '')
    
    success = DocumentRelationships.add_relationship(
        source_id,
        target_id,
        relationship_type,
        description
    )
    
    if success:
        flash('Powiązanie dodane pomyślnie!', 'success')
    else:
        flash('Nie udało się dodać powiązania', 'error')
    
    return redirect(url_for('collab_routes.view_relationships', document_id=source_id))

# =====================================
# NOTIFICATIONS ROUTES
# =====================================

@collab_routes.route('/notifications')
@login_required
def notifications_view():
    """View user notifications"""
    notifications = NotificationSystem.get_unread_notifications(current_user.id)
    
    return render_template(
        'collaboration/notifications.html',
        notifications=notifications
    )

@collab_routes.route('/notifications/unread')
@login_required
def get_unread_notifications():
    """Get unread notifications count (for AJAX)"""
    notifications = NotificationSystem.get_unread_notifications(current_user.id)
    
    return jsonify({
        'success': True,
        'count': len(notifications),
        'notifications': notifications[:5]  # Latest 5
    })

# =====================================
# VERSION COMPARISON
# =====================================

@collab_routes.route('/compare/<int:document_id>')
@login_required
def compare_versions(document_id):
    """Compare document versions"""
    from aircloud_legal_platform import Document, DocumentVersion
    
    document = Document.query.get_or_404(document_id)
    versions = DocumentVersion.query.filter_by(
        document_id=document_id
    ).order_by(DocumentVersion.version_number.desc()).all()
    
    version1_id = request.args.get('v1', type=int)
    version2_id = request.args.get('v2', type=int)
    
    diff_html = None
    if version1_id and version2_id:
        v1 = DocumentVersion.query.get(version1_id)
        v2 = DocumentVersion.query.get(version2_id)
        
        if v1 and v2:
            import difflib
            diff = difflib.unified_diff(
                v1.content.splitlines(),
                v2.content.splitlines(),
                lineterm=''
            )
            diff_html = '<br>'.join(diff)
    
    return render_template(
        'collaboration/version_compare.html',
        document=document,
        versions=versions,
        diff_html=diff_html
    )

# =====================================
# EXPORT ROUTES
# =====================================

@collab_routes.route('/export/<int:document_id>/<format>')
@login_required
def export_document(document_id, format):
    """Export document in various formats"""
    from aircloud_legal_platform import Document
    from flask import send_file
    import tempfile
    import os
    
    document = Document.query.get_or_404(document_id)
    
    if format == 'pdf':
        # Use ReportLab for PDF generation
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import simpleSplit
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        c = canvas.Canvas(temp_file.name, pagesize=A4)
        width, height = A4
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, document.title)
        
        # Content (simplified)
        c.setFont("Helvetica", 12)
        y = height - 100
        for line in document.content.split('\n')[:50]:  # First 50 lines
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(50, y, line[:80])  # Max 80 chars per line
            y -= 15
        
        c.save()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'{document.title}.pdf',
            mimetype='application/pdf'
        )
    
    elif format == 'docx':
        from docx import Document as DocxDocument
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        
        doc = DocxDocument()
        doc.add_heading(document.title, 0)
        doc.add_paragraph(document.content)
        doc.save(temp_file.name)
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'{document.title}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    else:
        flash('Nieobsługiwany format eksportu', 'error')
        return redirect(url_for('index'))

print("✅ Aircloud Collaboration Routes Loaded!")
